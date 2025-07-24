import os
import re
from bs4 import BeautifulSoup
from fpdf import FPDF
from docx import Document

def get_templates():
    return ["Modern", "Classic", "Minimal"]

def save_feedback(user, feedback):
    os.makedirs("feedback", exist_ok=True)
    with open(f"feedback/{user.replace(' ', '_').lower()}.txt", "a", encoding="utf-8") as f:
        f.write(feedback.strip() + "\n\n")

def generate_resume(name, email, phone, linkedin, github,
                    summary, skills, experience, education,
                    template="Modern", job_description="", theme="Light"):
    # Basic HTML formatting for demo
    html = f"""
    <html>
    <head>
    <style>
        body {{
            font-family: Arial, sans-serif;
            background-color: {"#fff" if theme == "Light" else "#1e1e1e"};
            color: {"#000" if theme == "Light" else "#fff"};
            padding: 20px;
        }}
    </style>
    </head>
    <body>
        <h1>{name}</h1>
        <p><strong>Email:</strong> {email} | <strong>Phone:</strong> {phone}</p>
        <p><strong>LinkedIn:</strong> {linkedin} | <strong>GitHub:</strong> {github}</p>
        <h2>Summary</h2>
        <p>{summary}</p>
        <h2>Skills</h2>
        <p>{skills}</p>
        <h2>Experience</h2>
        <p>{experience}</p>
        <h2>Education</h2>
        <p>{education}</p>
    </body>
    </html>
    """

    # ATS Score
    ats_score = 0
    if job_description:
        resume_text = " ".join([summary, skills, experience, education]).lower()
        job_keywords = set(re.findall(r'\b\w+\b', job_description.lower()))
        matched = [kw for kw in job_keywords if kw in resume_text]
        ats_score = (len(matched) / len(job_keywords)) * 100 if job_keywords else 0

    return html, ats_score

def html_to_pdf(html_path, pdf_path):
    try:
        with open(html_path, "r", encoding="utf-8") as f:
            html = f.read()

        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text()

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=12)
        for line in text.split("\n"):
            pdf.multi_cell(0, 10, line)
        pdf.output(pdf_path)
        return None
    except Exception as e:
        return str(e)

def html_to_docx(html_path, docx_path):
    with open(html_path, "r", encoding="utf-8") as f:
        html = f.read()

    soup = BeautifulSoup(html, "html.parser")
    doc = Document()
    doc.add_heading("Resume", 0)
    for element in soup.find_all(["h1", "h2", "p"]):
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        else:
            doc.add_paragraph(element.get_text())
    doc.save(docx_path)

def suggest_keywords(job_description, current_skills):
    job_description = job_description.lower()
    current_skills = current_skills.lower()

    roles = {
        "data scientist": ["python", "machine learning", "sql", "pandas", "numpy"],
        "web developer": ["html", "css", "javascript", "react", "node.js"],
        "data analyst": ["excel", "sql", "tableau", "powerbi", "statistics"],
        "ai engineer": ["deep learning", "tensorflow", "pytorch", "nlp", "mlops"],
        "software engineer": ["java", "c++", "python", "git", "linux"],
    }

    suggestions = []
    for role, keywords in roles.items():
        if role in job_description:
            for kw in keywords:
                if kw not in current_skills:
                    suggestions.append(kw)
            break  # only apply the first matching role

    return suggestions
